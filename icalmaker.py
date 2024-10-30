from docx import Document
from ics import Calendar, Event
from datetime import datetime, timedelta
import pytz
import os
import re
import subprocess  # To run Git commands

# Set the .docx file names and paths as variables
DOCX_FILE_NAME_1 = "Aut2.docx"  # Relative path in parent directory
DOCX_FILE_NAME_2 = r"C:\Users\Tom\Shirley Community Primary School\Shirley All Staff - General\2024-25 Termly Plan\2425 Shirley Termly Plan.docx"  # Full Windows path with spaces
MAX_HEADER_LENGTH = 15  # Set the maximum number of characters for fallback trimming

# Dictionary of key phrases to shorter names
header_mapping = {
    "Professional Development": "ProfessionalDevelopment",
    "Assessment Cycle": "AssessmentCycle",
    "Leadership/SLT": "SLT",
    "Monitoring Cycle": "MonitoringCycle",
    "School Events inc information": "SchoolEvents"
}

def trim_header(header):
    # Flatten the header to a single line for matching by removing line breaks
    flattened_header = " ".join(header.splitlines())
    
    # Check if any key phrase in header_mapping is found in the header text
    for key_phrase, short_header in header_mapping.items():
        if key_phrase in flattened_header:
            return short_header
    
    # If no specific match, use a generic transformation:
    # 1. Remove special characters, capitalise each word, and remove whitespace
    generic_trimmed = re.sub(r'[^A-Za-z0-9]', '', flattened_header.title())
    
    # 2. Trim to a max number of characters if it exceeds the set limit
    if len(generic_trimmed) > MAX_HEADER_LENGTH:
        generic_trimmed = generic_trimmed[:MAX_HEADER_LENGTH]
    
    return generic_trimmed

def generate_ics_from_docx(docx_file_name_1, docx_file_name_2, ics_output_name):
    # Define paths
    word_doc_path_1 = os.path.join("..", docx_file_name_1)  # Parent directory for Word doc
    word_doc_path_2 = docx_file_name_2  # Full path for second Word doc
    ics_output_path = ics_output_name  # Current directory for .ics file

    # Check for the existence of both documents
    doc1_exists = os.path.exists(word_doc_path_1)
    doc2_exists = os.path.exists(word_doc_path_2)

    if not doc1_exists and not doc2_exists:
        print("Error: Both DOCX files are missing.")
        return
    elif not doc1_exists:
        print(f"Warning: DOCX file '{word_doc_path_1}' is missing. Only '{word_doc_path_2}' will be processed.")
    elif not doc2_exists:
        print(f"Warning: DOCX file '{word_doc_path_2}' is missing. Only '{word_doc_path_1}' will be processed.")
    
    # Load documents if they exist
    calendar = Calendar()

    def process_doc(doc):
        headers = []
        week_data = []

        for table in doc.tables:
            # Extract headers and store in headers list
            headers = [cell.text.strip() for cell in table.rows[0].cells if cell.text.strip() != '']
            
            # Read each row following the headers
            for row in table.rows[1:]:
                week_row = {}
                week_beg_date_str = row.cells[0].text.strip()  # First cell for 'WeekBeg' date
                try:
                    # Attempt to parse the date in the first cell as WeekBeg
                    week_beg_date = datetime.strptime(week_beg_date_str, "%d.%m.%y")
                except ValueError:
                    continue  # Skip rows that do not contain a valid date in the first cell
                
                # Store each cell’s events in week_row with header as key
                for col_index, cell in enumerate(row.cells[1:], start=1):
                    if col_index < len(headers):
                        # Trim the header if necessary
                        header = trim_header(headers[col_index])
                        events = cell.text.strip().splitlines()
                        week_row[header] = [event.strip() for event in events if event.strip()]
                
                # Add data with start date of the week
                week_data.append((week_beg_date, week_row))

        # Create ICS events from extracted data
        for week_start, events_dict in week_data:
            # Define the start and end dates as Monday to Friday
            event_start_date = week_start  # Monday of the week
            event_end_date = week_start + timedelta(days=4)  # Friday of the week
            
            for header, events in events_dict.items():
                for event_description in events:
                    # Create an all-day event
                    event = Event()
                    event.name = f"{header} - {event_description}"
                    event.begin = event_start_date
                    event.end = event_end_date
                    event.make_all_day()  # Makes it an all-day event from Monday to Friday
                    event.description = f"Automatically generated task for {header}"
                    calendar.events.add(event)

    # Process each document if it exists
    if doc1_exists:
        process_doc(Document(word_doc_path_1))
    if doc2_exists:
        process_doc(Document(word_doc_path_2))
    
    # Write the ICS file
    with open(ics_output_path, 'w', encoding="utf-8") as f:
        f.writelines(calendar.serialize_iter())
    
    print(f"ICS file created at: {ics_output_path}")

    # Upload to GitHub
    upload_to_github(ics_output_path)

def upload_to_github(file_path):
    try:
        # Run Git commands to upload the .ics file to GitHub
        subprocess.run(["git", "add", file_path], check=True)
        subprocess.run(["git", "commit", "-m", "Add updated .ics calendar file"], check=True)
        subprocess.run(["git", "push"], check=True)
        print(f"{file_path} has been uploaded to GitHub.")
    except subprocess.CalledProcessError as e:
        print(f"An error occurred while uploading to GitHub: {e}")

# Example usage:
generate_ics_from_docx(DOCX_FILE_NAME_1, DOCX_FILE_NAME_2, "output.ics")
